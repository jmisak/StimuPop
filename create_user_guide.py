"""
StimuPop User Guide Generator

Creates a professional, graphically appealing DOCX user guide
for the StimuPop Excel to PowerPoint converter application.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc):
    """Add a horizontal line/divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Create bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_styled_heading(doc, text, level=1):
    """Create a styled heading with custom formatting."""
    heading = doc.add_heading(text, level=level)

    # Style based on level
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        heading.runs[0].font.size = Pt(24)
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        heading.runs[0].font.size = Pt(18)
    elif level == 3:
        heading.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        heading.runs[0].font.size = Pt(14)

    return heading


def add_info_box(doc, title, content, box_color="E7F3FF", border_color="2E74B5"):
    """Add a styled information box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Set cell background
    set_cell_shading(cell, box_color)

    # Set border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

    # Add title
    title_para = cell.paragraphs[0]
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # Add content
    content_para = cell.add_paragraph()
    content_run = content_para.add_run(content)
    content_run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing


def add_warning_box(doc, content):
    """Add a warning/caution box."""
    add_info_box(doc, "⚠️ Important", content, "FFF3CD", "856404")


def add_tip_box(doc, content):
    """Add a tip/hint box."""
    add_info_box(doc, "💡 Tip", content, "D4EDDA", "155724")


def add_step_table(doc, steps):
    """Add a numbered steps table with visual styling."""
    table = doc.add_table(rows=len(steps), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    table.columns[0].width = Inches(0.6)
    table.columns[1].width = Inches(5.4)

    for idx, (step_num, step_text) in enumerate(steps):
        # Number cell
        num_cell = table.cell(idx, 0)
        set_cell_shading(num_cell, "2E74B5")
        num_para = num_cell.paragraphs[0]
        num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_run = num_para.add_run(str(step_num))
        num_run.bold = True
        num_run.font.size = Pt(14)
        num_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Text cell
        text_cell = table.cell(idx, 1)
        if idx % 2 == 0:
            set_cell_shading(text_cell, "F8F9FA")
        text_para = text_cell.paragraphs[0]
        text_para.paragraph_format.left_indent = Pt(6)
        text_run = text_para.add_run(step_text)
        text_run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing


def create_user_guide():
    """Generate the complete user guide document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ========== TITLE PAGE ==========
    doc.add_paragraph()
    doc.add_paragraph()

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("🎯 StimuPop")
    title_run.bold = True
    title_run.font.size = Pt(48)
    title_run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Excel to PowerPoint Converter")
    sub_run.font.size = Pt(24)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # User Guide label
    guide_label = doc.add_paragraph()
    guide_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide_run = guide_label.add_run("USER GUIDE")
    guide_run.bold = True
    guide_run.font.size = Pt(18)
    guide_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    doc.add_paragraph()

    # Version info box
    version_table = doc.add_table(rows=1, cols=1)
    version_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    version_cell = version_table.cell(0, 0)
    set_cell_shading(version_cell, "E7F3FF")
    version_para = version_cell.paragraphs[0]
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run("Version 7.0.0")
    version_run.font.size = Pt(14)
    version_run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # Page break
    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    create_styled_heading(doc, "Table of Contents", 1)
    add_horizontal_line(doc)

    toc_items = [
        ("1.", "Introduction", "What is StimuPop?"),
        ("2.", "Getting Started", "Installation and first launch"),
        ("3.", "Preparing Your Excel File", "Data structure requirements"),
        ("4.", "Using StimuPop", "Step-by-step guide"),
        ("5.", "Configuration Options", "Customizing your output"),
        ("6.", "Troubleshooting", "Common issues and solutions"),
        ("7.", "Quick Reference", "Keyboard shortcuts and tips"),
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    for idx, (num, title, desc) in enumerate(toc_items):
        toc_table.cell(idx, 0).paragraphs[0].add_run(num).bold = True
        toc_table.cell(idx, 1).paragraphs[0].add_run(title).bold = True
        toc_table.cell(idx, 2).paragraphs[0].add_run(desc).font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ========== SECTION 1: INTRODUCTION ==========
    create_styled_heading(doc, "1. Introduction", 1)
    add_horizontal_line(doc)

    create_styled_heading(doc, "What is StimuPop?", 2)

    intro_para = doc.add_paragraph()
    intro_para.add_run("StimuPop").bold = True
    intro_para.add_run(" is a powerful yet easy-to-use application that converts Excel spreadsheet data into professional PowerPoint presentations. It's designed to save you hours of manual work by automatically:")

    features = [
        "Extracting embedded images from Excel cells",
        "Creating individual slides for each data row",
        "Uniform image sizing with multiple sizing modes",
        "Formatting text with custom fonts, sizes, and colors",
        "Supporting both portrait and landscape orientations",
        "Handling errors gracefully (missing images won't stop the process)"
    ]

    for feature in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run("✓ ").font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
        p.add_run(feature)

    doc.add_paragraph()
    add_tip_box(doc, "StimuPop is perfect for creating product catalogs, photo albums, training materials, and any presentation where each slide follows the same format.")

    # ========== SECTION 2: GETTING STARTED ==========
    doc.add_page_break()
    create_styled_heading(doc, "2. Getting Started", 1)
    add_horizontal_line(doc)

    create_styled_heading(doc, "Installation (Portable Version)", 2)

    doc.add_paragraph("The portable version requires no installation. Simply follow these steps:")

    install_steps = [
        ("1", "Extract the ZIP file to any folder on your computer"),
        ("2", "Open the extracted folder"),
        ("3", "Double-click 'StimuPop.exe' (or 'StimuPop.bat') to launch the application"),
        ("4", "Wait for the terminal window to display 'Server ready!' message"),
        ("5", "Hold Ctrl and click the localhost link in the terminal, OR copy the URL and paste it into your browser"),
    ]
    add_step_table(doc, install_steps)

    add_info_box(doc, "📌 First Launch",
                 "The first time you run StimuPop, the server may take 30-60 seconds to start. "
                 "Wait for the 'Server ready!' message before clicking the link. "
                 "Subsequent launches will be faster.")

    create_styled_heading(doc, "System Requirements", 2)

    req_table = doc.add_table(rows=4, cols=2)
    req_table.style = 'Table Grid'
    requirements = [
        ("Operating System", "Windows 10 or later (64-bit)"),
        ("Disk Space", "~300 MB for portable version"),
        ("Memory", "4 GB RAM minimum"),
        ("Browser", "Chrome, Firefox, or Edge (latest version)"),
    ]
    for idx, (req, val) in enumerate(requirements):
        req_table.cell(idx, 0).paragraphs[0].add_run(req).bold = True
        set_cell_shading(req_table.cell(idx, 0), "F0F0F0")
        req_table.cell(idx, 1).paragraphs[0].add_run(val)

    doc.add_paragraph()

    # ========== SECTION 3: PREPARING YOUR EXCEL FILE ==========
    doc.add_page_break()
    create_styled_heading(doc, "3. Preparing Your Excel File", 1)
    add_horizontal_line(doc)

    doc.add_paragraph("For best results, structure your Excel file as follows:")

    create_styled_heading(doc, "Recommended Structure", 2)

    # Example table
    example_table = doc.add_table(rows=4, cols=4)
    example_table.style = 'Table Grid'

    headers = ["Column A\n(Skip)", "Column B\n(Image)", "Column C\n(Title)", "Column D\n(Description)"]
    for idx, header in enumerate(headers):
        cell = example_table.cell(0, idx)
        set_cell_shading(cell, "2E74B5")
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    data_rows = [
        ["ID-001", "[Image 1]", "Product One", "Description of product one"],
        ["ID-002", "[Image 2]", "Product Two", "Description of product two"],
        ["ID-003", "[Image 3]", "Product Three", "Description of product three"],
    ]

    for row_idx, row_data in enumerate(data_rows, start=1):
        for col_idx, cell_data in enumerate(row_data):
            cell = example_table.cell(row_idx, col_idx)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F8F9FA")
            cell.paragraphs[0].add_run(cell_data).font.size = Pt(10)

    doc.add_paragraph()

    create_styled_heading(doc, "Image Options", 2)

    doc.add_paragraph("StimuPop supports three ways to include images:")

    img_options = [
        ("Embedded Images", "Paste images directly into Excel cells. This is the most reliable method."),
        ("File Paths", "Enter the full path to an image file (e.g., C:\\Images\\photo.jpg)"),
        ("URLs", "Enter a web URL to an image (e.g., https://example.com/image.png)"),
    ]

    for title, desc in img_options:
        p = doc.add_paragraph()
        p.add_run(f"• {title}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph()
    add_warning_box(doc, "Images embedded directly in Excel cells provide the most consistent results. "
                        "File paths must be accessible from the computer running StimuPop.")

    # Column Flexibility note (Item #19)
    create_styled_heading(doc, "Column Flexibility", 2)

    doc.add_paragraph("Your Excel file can have any number of columns. StimuPop is flexible about your data structure:")

    flex_points = [
        "Add as many text columns as you need (C, D, E, F, G, etc.)",
        "Skip columns you do not want to include in the presentation",
        "Use column letters (B, C, D) or column header names (Image, Title, Description)",
        "Columns can be in any order - just specify the correct references",
        "Empty cells are handled gracefully (no errors for missing data)",
    ]

    for point in flex_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(point)

    doc.add_paragraph()
    add_tip_box(doc, "You only need to specify the columns you want to use. StimuPop ignores all other columns in your Excel file.")

    # ========== SECTION 4: USING STIMUPOP ==========
    doc.add_page_break()
    create_styled_heading(doc, "4. Using StimuPop", 1)
    add_horizontal_line(doc)

    create_styled_heading(doc, "Step-by-Step Guide", 2)

    usage_steps = [
        ("1", "Launch StimuPop by double-clicking 'StimuPop.exe' and wait for 'Server ready!'"),
        ("2", "Ctrl+click the localhost link (or copy/paste into browser)"),
        ("3", "Click 'Browse files' under 'Upload Excel File' and select your .xlsx file"),
        ("4", "Optionally upload a PowerPoint template for custom styling"),
        ("5", "Set the Image Column (e.g., 'B' or the column header name)"),
        ("6", "Set the Text Columns (e.g., 'C,D,E,F' for multiple columns)"),
        ("7", "Adjust font size using the slider (default: 14pt)"),
        ("8", "Click the blue 'Generate Presentation' button"),
        ("9", "Wait for processing (progress bar shows status)"),
        ("10", "Click 'Download Presentation' to save your .pptx file"),
    ]
    add_step_table(doc, usage_steps)

    create_styled_heading(doc, "Understanding the Interface", 2)

    interface_desc = doc.add_paragraph()
    interface_desc.add_run("The StimuPop interface is divided into several sections:\n\n")

    sections = [
        ("Upload Files", "Where you select your Excel file and optional PowerPoint template"),
        ("Configuration", "Basic settings like image column, text columns, and font size"),
        ("Advanced Settings", "Layout options and per-column formatting (click to expand)"),
        ("Data Preview", "Shows a preview of your Excel data before generation"),
        ("Generate Button", "Starts the presentation generation process"),
    ]

    for section, desc in sections:
        p = doc.add_paragraph()
        p.add_run(f"📍 {section}: ").bold = True
        p.add_run(desc)

    # ========== SECTION 5: CONFIGURATION OPTIONS ==========
    doc.add_page_break()
    create_styled_heading(doc, "5. Configuration Options", 1)
    add_horizontal_line(doc)

    create_styled_heading(doc, "Basic Settings", 2)

    basic_settings = [
        ("Image Column", "Letter (A-Z) or name of the column containing images", "B"),
        ("Text Columns", "Comma-separated list of columns for text content", "C,D,E,F"),
        ("Font Size", "Default font size for all text (8-48pt)", "14"),
        ("Pictures Only", "Skip all text columns and create image-only slides", "Off"),
    ]

    settings_table = doc.add_table(rows=len(basic_settings)+1, cols=3)
    settings_table.style = 'Table Grid'

    # Header
    for idx, header in enumerate(["Setting", "Description", "Default"]):
        cell = settings_table.cell(0, idx)
        set_cell_shading(cell, "2E74B5")
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, (setting, desc, default) in enumerate(basic_settings, start=1):
        settings_table.cell(row_idx, 0).paragraphs[0].add_run(setting).bold = True
        settings_table.cell(row_idx, 1).paragraphs[0].add_run(desc)
        settings_table.cell(row_idx, 2).paragraphs[0].add_run(default)

    doc.add_paragraph()

    # Pictures Only Mode (added in v6.2)
    create_styled_heading(doc, "Pictures Only Mode", 3)

    doc.add_paragraph("Enable 'Pictures Only (no text)' checkbox to create image-only slideshows:")

    pictures_only_uses = [
        "Photo albums and image galleries",
        "Product image catalogs without descriptions",
        "Visual presentations where text is unnecessary",
        "Quick image slideshows from Excel collections",
    ]

    for use in pictures_only_uses:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(use)

    doc.add_paragraph()
    add_tip_box(doc, "When Pictures Only mode is enabled, the Text Columns setting is ignored. Only the Image Column is used to create slides.")

    create_styled_heading(doc, "Advanced Settings", 2)

    doc.add_paragraph("Click 'Advanced Settings' to access template mode, layout, image sizing, text spacing, and alignment options:")

    # Template Mode Summary (Item #15)
    create_styled_heading(doc, "Template Mode", 3)

    doc.add_paragraph("StimuPop offers two generation modes for creating slides:")

    template_modes = [
        ("Blank Mode (default)", "Creates slides from scratch with configurable layout. You control image position, text position, fonts, and colors. Best for new presentations or when you need full control over styling."),
        ("Template Mode", "Uses an existing PowerPoint template slide as the base. StimuPop clones the template for each row and populates placeholders with your data. Best when you have an existing design you want to preserve exactly."),
    ]

    for mode, desc in template_modes:
        p = doc.add_paragraph()
        p.add_run(f"{mode}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph()

    # v7.0 Template Mode Enhancement
    create_styled_heading(doc, "Template Mode Enhancement (v7.0)", 3)

    doc.add_paragraph("Version 7.0 introduces dynamic column mapping for Template Mode:")

    template_v7_features = [
        "Template Mode now automatically maps your Excel columns to template paragraphs",
        "Works with any column letters (A, B, C or custom names)",
        "Empty paragraphs in your template are preserved as spacers",
        "No longer limited to specific column configurations",
    ]

    for feature in template_v7_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(feature)

    doc.add_paragraph()
    add_tip_box(doc, "v7.0 Enhancement: Dynamic column mapping - your columns are automatically matched to non-empty paragraphs in your template, regardless of which column letters you use. This means you can use columns B and C, or D and E, or any combination that fits your data structure.")

    doc.add_paragraph()
    add_info_box(doc, "Template Mode Settings",
                 "When using Template Mode:\n"
                 "- Upload a .pptx template with your desired slide design\n"
                 "- Specify the Image Placeholder Name (shape containing the image)\n"
                 "- Specify the Text Placeholder Name (shape containing text)\n"
                 "- StimuPop preserves all template formatting (fonts, sizes, colors)\n"
                 "- Your columns are dynamically mapped to non-empty template paragraphs")

    # Image Sizing subsection
    create_styled_heading(doc, "Image Sizing", 3)

    doc.add_paragraph("Control how images are sized uniformly across all slides. These settings work in both Blank and Template modes:")

    # Size modes table
    size_modes_table = doc.add_table(rows=5, cols=2)
    size_modes_table.style = 'Table Grid'

    size_modes = [
        ("Size Mode", "Description"),
        ("Fit to Box (Recommended)", "Images scale to fit within Max Width and Max Height while maintaining aspect ratio. All images appear uniform."),
        ("Fit Width Only", "Images have fixed width, height adjusts automatically based on aspect ratio."),
        ("Fit Height Only", "Images have fixed height, width adjusts automatically based on aspect ratio."),
        ("Stretch to Exact Size", "Images are forced to exact dimensions. May distort images."),
    ]

    for idx, (mode, desc) in enumerate(size_modes):
        if idx == 0:
            set_cell_shading(size_modes_table.cell(idx, 0), "2E74B5")
            set_cell_shading(size_modes_table.cell(idx, 1), "2E74B5")
            size_modes_table.cell(idx, 0).paragraphs[0].add_run(mode).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            size_modes_table.cell(idx, 1).paragraphs[0].add_run(desc).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            size_modes_table.cell(idx, 0).paragraphs[0].runs[0].bold = True
            size_modes_table.cell(idx, 1).paragraphs[0].runs[0].bold = True
        else:
            size_modes_table.cell(idx, 0).paragraphs[0].add_run(mode).bold = True
            size_modes_table.cell(idx, 1).paragraphs[0].add_run(desc)

    doc.add_paragraph()

    doc.add_paragraph("Image dimension sliders (updated in v6.2):")
    img_sliders = [
        ("Max Width", "0.0 - 9.0 inches (0 allows thumbnail-sized images)"),
        ("Max Height", "0.0 - 7.0 inches (0 allows thumbnail-sized images)"),
    ]
    for slider, range_desc in img_sliders:
        p = doc.add_paragraph()
        p.add_run(f"- {slider}: ").bold = True
        p.add_run(range_desc)

    doc.add_paragraph()
    add_tip_box(doc, "Use 'Fit to Box' mode for product catalogs and photo albums where you want all images to appear the same size regardless of their original dimensions.")

    # Layout settings
    create_styled_heading(doc, "Layout Position (Blank Mode Only)", 3)

    doc.add_paragraph("These settings control slide layout and are only available when using Blank slide mode (not Template mode):")

    advanced_settings = [
        ("Image Top Position", "Distance from top of slide to image", "0.5 inches"),
        ("Text Top Position", "Distance from top of slide to text area (0.0-9.0 inches)", "5.0 inches"),
        ("Slide Orientation", "Portrait (tall) or Landscape (wide)", "Portrait"),
    ]

    for setting, desc, default in advanced_settings:
        p = doc.add_paragraph()
        p.add_run(f"- {setting}: ").bold = True
        p.add_run(f"{desc} (Default: {default})")

    doc.add_paragraph()

    # Image Alignment subsection (Moved after Layout Position per item #16)
    create_styled_heading(doc, "Image Alignment (Blank Mode Only)", 3)

    doc.add_paragraph("Control how images are positioned within their bounding box. These settings only apply to Blank mode:")

    # Alignment options table
    align_table = doc.add_table(rows=3, cols=2)
    align_table.style = 'Table Grid'

    align_options = [
        ("Alignment Type", "Options"),
        ("Vertical Alignment", "Top, Center (default), Bottom"),
        ("Horizontal Alignment", "Left, Center (default), Right"),
    ]

    for idx, (opt_type, options) in enumerate(align_options):
        if idx == 0:
            set_cell_shading(align_table.cell(idx, 0), "2E74B5")
            set_cell_shading(align_table.cell(idx, 1), "2E74B5")
            align_table.cell(idx, 0).paragraphs[0].add_run(opt_type).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            align_table.cell(idx, 1).paragraphs[0].add_run(options).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            align_table.cell(idx, 0).paragraphs[0].runs[0].bold = True
            align_table.cell(idx, 1).paragraphs[0].runs[0].bold = True
        else:
            align_table.cell(idx, 0).paragraphs[0].add_run(opt_type).bold = True
            align_table.cell(idx, 1).paragraphs[0].add_run(options)

    doc.add_paragraph()
    add_tip_box(doc, "Use 'Bottom' vertical alignment for variety cards where you want images anchored to the bottom of the image area, regardless of image height.")

    # Text Spacing subsection (Item #17 and #21)
    create_styled_heading(doc, "Text Spacing", 3)

    doc.add_paragraph("Control spacing and overflow behavior for text content:")

    text_spacing_settings = [
        ("Paragraph Spacing", "Space after each paragraph (0-24pt). Default 0pt means no extra spacing between lines."),
        ("Text Overflow", "How text boxes handle content that exceeds the box size:"),
    ]

    for setting, desc in text_spacing_settings:
        p = doc.add_paragraph()
        p.add_run(f"- {setting}: ").bold = True
        p.add_run(desc)

    # Text Overflow options (NEW in v6.2)
    doc.add_paragraph()
    overflow_options = [
        ("Resize shape to fit text (default)", "The text box expands to fit all content. Text remains at original size."),
        ("Shrink text on overflow", "Text automatically shrinks to fit within the text box boundaries."),
    ]

    for opt, desc in overflow_options:
        p = doc.add_paragraph()
        p.add_run(f"    - {opt}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph()
    add_info_box(doc, "Text Overflow",
                 "The Text Overflow dropdown gives you control over how StimuPop handles long text that exceeds the text box. "
                 "'Resize shape to fit text' keeps text readable but may extend beyond the slide. "
                 "'Shrink text on overflow' keeps text within bounds but may make it smaller.")

    # Advanced Positioning (added in v6.0)
    create_styled_heading(doc, "Advanced Positioning", 3)

    doc.add_paragraph("For precise control over text column placement, enable 'Advanced Positioning' by checking the checkbox in Advanced Settings. This reveals per-column position controls:")

    position_modes = [
        ("Auto (default)", "Text flows sequentially after the previous column. Position depends on content length of earlier columns."),
        ("Fixed", "Text is placed at an exact position on the slide, regardless of other content. Creates a separate text box."),
    ]

    for mode, desc in position_modes:
        p = doc.add_paragraph()
        p.add_run(f"• {mode}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph("Default fixed positions (useful for variety cards):")

    default_positions = [
        ("Column E", "5.0 inches from top"),
        ("Column F", "6.5 inches from top"),
    ]

    for col, pos in default_positions:
        p = doc.add_paragraph()
        p.add_run(f"• {col}: ").bold = True
        p.add_run(pos)

    doc.add_paragraph()
    add_tip_box(doc, "Use Fixed positioning for columns E and F when you want them to appear in the same location on every slide, even if the Brand (C) and Product Name (D) vary in length.")

    create_styled_heading(doc, "Per-Column Formatting (Blank Mode Only)", 3)

    doc.add_paragraph("When using Blank mode, you can customize each text column individually in Advanced Settings:")

    format_options = ["Font size (8-48pt, expanded range in v6.2)", "Font family (Calibri, Arial, Times New Roman, etc.)",
                      "Text color (color picker)", "Bold and Italic styles"]

    for opt in format_options:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(opt)

    doc.add_paragraph()
    add_tip_box(doc, "Use larger, bold fonts for titles (Column C) and smaller regular fonts for descriptions (Column D) to create visual hierarchy.")

    # ========== SECTION 6: TROUBLESHOOTING ==========
    doc.add_page_break()
    create_styled_heading(doc, "6. Troubleshooting", 1)
    add_horizontal_line(doc)

    issues = [
        ("Application won't start",
         "• Ensure you extracted all files from the ZIP\n"
         "• Run as Administrator (right-click → Run as administrator)\n"
         "• Check that antivirus isn't blocking the application"),

        ("Browser doesn't open automatically",
         "• Wait for 'Server ready!' message in the terminal window\n"
         "• Hold Ctrl and click the localhost link, or copy/paste the URL\n"
         "• Manually open your browser and go to: http://localhost:8501\n"
         "• Try a different browser (Chrome recommended)"),

        ("'Column not found' error",
         "• Check that your column letter/name matches exactly\n"
         "• Column letters are case-insensitive (B = b)\n"
         "• Column names ARE case-sensitive"),

        ("Images not appearing on slides",
         "• Ensure images are embedded in cells, not floating\n"
         "• Check file paths are correct and accessible\n"
         "• Verify image files are JPG, PNG, or GIF format"),

        ("Presentation generation is slow",
         "• Large images take longer to process\n"
         "• Consider resizing images in Excel before upload\n"
         "• Reduce the number of slides if testing"),

        ("Text formatting not applied",
         "• Ensure you configured formatting in Advanced Settings\n"
         "• Check that column letters in formatting match your data"),

        ("Images are different sizes on slides",
         "• Open Advanced Settings and set Size Mode to 'Fit to Box'\n"
         "• Adjust Max Width and Max Height to your desired dimensions\n"
         "• All images will now scale uniformly"),
    ]

    for issue, solution in issues:
        p = doc.add_paragraph()
        p.add_run(f"❓ {issue}").bold = True
        p.add_run(f"\n{solution}")
        doc.add_paragraph()

    # ========== SECTION 7: QUICK REFERENCE ==========
    doc.add_page_break()
    create_styled_heading(doc, "7. Quick Reference", 1)
    add_horizontal_line(doc)

    create_styled_heading(doc, "Column Reference Cheat Sheet", 2)

    ref_table = doc.add_table(rows=6, cols=2)
    ref_table.style = 'Table Grid'

    ref_data = [
        ("Reference Type", "Example"),
        ("Single letter", "B"),
        ("Multiple letters", "C,D,E,F"),
        ("Column name", "Title"),
        ("Mixed", "B,Title,Description"),
        ("With spaces", "\"Product Name\""),
    ]

    for idx, (col1, col2) in enumerate(ref_data):
        if idx == 0:
            set_cell_shading(ref_table.cell(idx, 0), "2E74B5")
            set_cell_shading(ref_table.cell(idx, 1), "2E74B5")
            ref_table.cell(idx, 0).paragraphs[0].add_run(col1).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            ref_table.cell(idx, 1).paragraphs[0].add_run(col2).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            ref_table.cell(idx, 0).paragraphs[0].runs[0].bold = True
            ref_table.cell(idx, 1).paragraphs[0].runs[0].bold = True
        else:
            ref_table.cell(idx, 0).paragraphs[0].add_run(col1)
            code_run = ref_table.cell(idx, 1).paragraphs[0].add_run(col2)
            code_run.font.name = 'Consolas'

    doc.add_paragraph()

    create_styled_heading(doc, "Keyboard Shortcuts", 2)

    shortcuts = [
        ("Ctrl + C", "Stop the application (in command window)"),
        ("F5", "Refresh the browser page"),
        ("Ctrl + S", "Save downloaded file (in browser)"),
    ]

    for shortcut, action in shortcuts:
        p = doc.add_paragraph()
        shortcut_run = p.add_run(f"  {shortcut}  ")
        shortcut_run.bold = True
        shortcut_run.font.name = 'Consolas'
        p.add_run(f"  →  {action}")

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer box
    add_info_box(doc, "📧 Need Help?",
                 "If you encounter issues not covered in this guide, please contact your system administrator "
                 "or the development team with:\n"
                 "• A description of the problem\n"
                 "• The exact error message (if any)\n"
                 "• Your Excel file (if possible)")

    # ========== FINAL PAGE ==========
    doc.add_page_break()

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    thanks = doc.add_paragraph()
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thanks_run = thanks.add_run("Thank you for using StimuPop!")
    thanks_run.font.size = Pt(20)
    thanks_run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    thanks_run.bold = True

    doc.add_paragraph()

    version_final = doc.add_paragraph()
    version_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_final.add_run("Version 7.0.0").font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Save document
    doc.save('StimuPop_User_Guide.docx')
    print("[OK] User Guide created successfully: StimuPop_User_Guide.docx")


def create_html_user_guide():
    """Generate the HTML version of the user guide."""
    html_content = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>StimuPop User Guide - v7.0.0</title>
    <style>
        :root {
            --primary-blue: #2E74B5;
            --secondary-blue: #4472C4;
            --text-gray: #666666;
            --light-bg: #F8F9FA;
            --info-bg: #E7F3FF;
            --warning-bg: #FFF3CD;
            --tip-bg: #D4EDDA;
            --warning-border: #856404;
            --tip-border: #155724;
        }
        * { box-sizing: border-box; }
        body {
            font-family: 'Segoe UI', Calibri, Arial, sans-serif;
            font-size: 16px;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: #fff;
        }
        h1 { color: var(--primary-blue); font-size: 2em; border-bottom: 2px solid var(--secondary-blue); padding-bottom: 10px; }
        h2 { color: var(--primary-blue); font-size: 1.5em; margin-top: 2em; }
        h3 { color: var(--secondary-blue); font-size: 1.2em; margin-top: 1.5em; }
        .title-page { text-align: center; padding: 60px 20px; }
        .title-page h1 { font-size: 3em; border: none; }
        .title-page .subtitle { font-size: 1.5em; color: var(--text-gray); }
        .title-page .version { display: inline-block; background: var(--info-bg); color: var(--primary-blue); padding: 10px 30px; border-radius: 5px; margin-top: 20px; }
        .info-box, .warning-box, .tip-box {
            padding: 15px 20px;
            border-radius: 5px;
            margin: 20px 0;
            border-left: 4px solid;
        }
        .info-box { background: var(--info-bg); border-color: var(--primary-blue); }
        .warning-box { background: var(--warning-bg); border-color: var(--warning-border); }
        .tip-box { background: var(--tip-bg); border-color: var(--tip-border); }
        .info-box strong, .warning-box strong, .tip-box strong { display: block; margin-bottom: 5px; }
        table { border-collapse: collapse; width: 100%; margin: 20px 0; }
        th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
        th { background: var(--primary-blue); color: white; }
        tr:nth-child(even) { background: var(--light-bg); }
        .step-table { margin: 20px 0; }
        .step-table .step-num {
            background: var(--primary-blue);
            color: white;
            font-weight: bold;
            text-align: center;
            width: 50px;
        }
        ul.features li::marker { color: #28A745; }
        code { font-family: Consolas, monospace; background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }
        .toc { background: var(--light-bg); padding: 20px; border-radius: 5px; }
        .toc ul { list-style: none; padding-left: 0; }
        .toc li { padding: 5px 0; }
        .toc a { color: var(--primary-blue); text-decoration: none; }
        .toc a:hover { text-decoration: underline; }
        @media print {
            body { max-width: none; }
            .title-page { page-break-after: always; }
            h1 { page-break-before: always; }
        }
    </style>
</head>
<body>

<div class="title-page">
    <h1>StimuPop</h1>
    <p class="subtitle">Excel to PowerPoint Converter</p>
    <p style="font-size: 1.2em; color: var(--secondary-blue); font-weight: bold;">USER GUIDE</p>
    <p class="version">Version 7.0.0</p>
</div>

<div class="toc">
    <h2 style="margin-top: 0;">Table of Contents</h2>
    <ul>
        <li><a href="#introduction">1. Introduction</a> - What is StimuPop?</li>
        <li><a href="#getting-started">2. Getting Started</a> - Installation and first launch</li>
        <li><a href="#preparing-excel">3. Preparing Your Excel File</a> - Data structure requirements</li>
        <li><a href="#using-stimupop">4. Using StimuPop</a> - Step-by-step guide</li>
        <li><a href="#configuration">5. Configuration Options</a> - Customizing your output</li>
        <li><a href="#troubleshooting">6. Troubleshooting</a> - Common issues and solutions</li>
        <li><a href="#quick-reference">7. Quick Reference</a> - Keyboard shortcuts and tips</li>
    </ul>
</div>

<h1 id="introduction">1. Introduction</h1>

<h2>What is StimuPop?</h2>
<p><strong>StimuPop</strong> is a powerful yet easy-to-use application that converts Excel spreadsheet data into professional PowerPoint presentations. It's designed to save you hours of manual work by automatically:</p>
<ul class="features">
    <li>Extracting embedded images from Excel cells</li>
    <li>Creating individual slides for each data row</li>
    <li>Uniform image sizing with multiple sizing modes</li>
    <li>Formatting text with custom fonts, sizes, and colors</li>
    <li>Supporting both portrait and landscape orientations</li>
    <li>Handling errors gracefully (missing images won't stop the process)</li>
</ul>
<div class="tip-box">
    <strong>Tip</strong>
    StimuPop is perfect for creating product catalogs, photo albums, training materials, and any presentation where each slide follows the same format.
</div>

<h1 id="getting-started">2. Getting Started</h1>

<h2>Installation (Portable Version)</h2>
<p>The portable version requires no installation. Simply follow these steps:</p>
<table class="step-table">
    <tr><td class="step-num">1</td><td>Extract the ZIP file to any folder on your computer</td></tr>
    <tr><td class="step-num">2</td><td>Open the extracted folder</td></tr>
    <tr><td class="step-num">3</td><td>Double-click 'StimuPop.exe' (or 'StimuPop.bat') to launch the application</td></tr>
    <tr><td class="step-num">4</td><td>Wait for the terminal window to display 'Server ready!' message</td></tr>
    <tr><td class="step-num">5</td><td>Hold Ctrl and click the localhost link in the terminal, OR copy the URL and paste it into your browser</td></tr>
</table>
<div class="info-box">
    <strong>First Launch</strong>
    The first time you run StimuPop, the server may take 30-60 seconds to start. Wait for the 'Server ready!' message before clicking the link. Subsequent launches will be faster.
</div>

<h2>System Requirements</h2>
<table>
    <tr><th>Requirement</th><th>Specification</th></tr>
    <tr><td><strong>Operating System</strong></td><td>Windows 10 or later (64-bit)</td></tr>
    <tr><td><strong>Disk Space</strong></td><td>~300 MB for portable version</td></tr>
    <tr><td><strong>Memory</strong></td><td>4 GB RAM minimum</td></tr>
    <tr><td><strong>Browser</strong></td><td>Chrome, Firefox, or Edge (latest version)</td></tr>
</table>

<h1 id="preparing-excel">3. Preparing Your Excel File</h1>
<p>For best results, structure your Excel file as follows:</p>

<h2>Recommended Structure</h2>
<table>
    <tr><th>Column A<br>(Skip)</th><th>Column B<br>(Image)</th><th>Column C<br>(Title)</th><th>Column D<br>(Description)</th></tr>
    <tr><td>ID-001</td><td>[Image 1]</td><td>Product One</td><td>Description of product one</td></tr>
    <tr><td>ID-002</td><td>[Image 2]</td><td>Product Two</td><td>Description of product two</td></tr>
    <tr><td>ID-003</td><td>[Image 3]</td><td>Product Three</td><td>Description of product three</td></tr>
</table>

<h2>Image Options</h2>
<p>StimuPop supports three ways to include images:</p>
<ul>
    <li><strong>Embedded Images:</strong> Paste images directly into Excel cells. This is the most reliable method.</li>
    <li><strong>File Paths:</strong> Enter the full path to an image file (e.g., C:\\Images\\photo.jpg)</li>
    <li><strong>URLs:</strong> Enter a web URL to an image (e.g., https://example.com/image.png)</li>
</ul>
<div class="warning-box">
    <strong>Important</strong>
    Images embedded directly in Excel cells provide the most consistent results. File paths must be accessible from the computer running StimuPop.
</div>

<h2>Column Flexibility</h2>
<p>Your Excel file can have any number of columns. StimuPop is flexible about your data structure:</p>
<ul>
    <li>Add as many text columns as you need (C, D, E, F, G, etc.)</li>
    <li>Skip columns you do not want to include in the presentation</li>
    <li>Use column letters (B, C, D) or column header names (Image, Title, Description)</li>
    <li>Columns can be in any order - just specify the correct references</li>
    <li>Empty cells are handled gracefully (no errors for missing data)</li>
</ul>
<div class="tip-box">
    <strong>Tip</strong>
    You only need to specify the columns you want to use. StimuPop ignores all other columns in your Excel file.
</div>

<h1 id="using-stimupop">4. Using StimuPop</h1>

<h2>Step-by-Step Guide</h2>
<table class="step-table">
    <tr><td class="step-num">1</td><td>Launch StimuPop by double-clicking 'StimuPop.exe' and wait for 'Server ready!'</td></tr>
    <tr><td class="step-num">2</td><td>Ctrl+click the localhost link (or copy/paste into browser)</td></tr>
    <tr><td class="step-num">3</td><td>Click 'Browse files' under 'Upload Excel File' and select your .xlsx file</td></tr>
    <tr><td class="step-num">4</td><td>Optionally upload a PowerPoint template for custom styling</td></tr>
    <tr><td class="step-num">5</td><td>Set the Image Column (e.g., 'B' or the column header name)</td></tr>
    <tr><td class="step-num">6</td><td>Set the Text Columns (e.g., 'C,D,E,F' for multiple columns)</td></tr>
    <tr><td class="step-num">7</td><td>Adjust font size using the slider (default: 14pt)</td></tr>
    <tr><td class="step-num">8</td><td>Click the blue 'Generate Presentation' button</td></tr>
    <tr><td class="step-num">9</td><td>Wait for processing (progress bar shows status)</td></tr>
    <tr><td class="step-num">10</td><td>Click 'Download Presentation' to save your .pptx file</td></tr>
</table>

<h2>Understanding the Interface</h2>
<p>The StimuPop interface is divided into several sections:</p>
<ul>
    <li><strong>Upload Files:</strong> Where you select your Excel file and optional PowerPoint template</li>
    <li><strong>Configuration:</strong> Basic settings like image column, text columns, and font size</li>
    <li><strong>Advanced Settings:</strong> Layout options and per-column formatting (click to expand)</li>
    <li><strong>Data Preview:</strong> Shows a preview of your Excel data before generation</li>
    <li><strong>Generate Button:</strong> Starts the presentation generation process</li>
</ul>

<h1 id="configuration">5. Configuration Options</h1>

<h2>Basic Settings</h2>
<table>
    <tr><th>Setting</th><th>Description</th><th>Default</th></tr>
    <tr><td><strong>Image Column</strong></td><td>Letter (A-Z) or name of the column containing images</td><td>B</td></tr>
    <tr><td><strong>Text Columns</strong></td><td>Comma-separated list of columns for text content</td><td>C,D,E,F</td></tr>
    <tr><td><strong>Font Size</strong></td><td>Default font size for all text (8-48pt)</td><td>14</td></tr>
    <tr><td><strong>Pictures Only</strong></td><td>Skip all text columns and create image-only slides</td><td>Off</td></tr>
</table>

<h3>Pictures Only Mode</h3>
<p>Enable 'Pictures Only (no text)' checkbox to create image-only slideshows:</p>
<ul>
    <li>Photo albums and image galleries</li>
    <li>Product image catalogs without descriptions</li>
    <li>Visual presentations where text is unnecessary</li>
    <li>Quick image slideshows from Excel collections</li>
</ul>
<div class="tip-box">
    <strong>Tip</strong>
    When Pictures Only mode is enabled, the Text Columns setting is ignored. Only the Image Column is used to create slides.
</div>

<h2>Advanced Settings</h2>
<p>Click 'Advanced Settings' to access template mode, layout, image sizing, text spacing, and alignment options:</p>

<h3>Template Mode</h3>
<p>StimuPop offers two generation modes for creating slides:</p>
<ul>
    <li><strong>Blank Mode (default):</strong> Creates slides from scratch with configurable layout. You control image position, text position, fonts, and colors. Best for new presentations or when you need full control over styling.</li>
    <li><strong>Template Mode:</strong> Uses an existing PowerPoint template slide as the base. StimuPop clones the template for each row and populates placeholders with your data. Best when you have an existing design you want to preserve exactly.</li>
</ul>

<h3>Template Mode Enhancement (v7.0)</h3>
<p>Version 7.0 introduces dynamic column mapping for Template Mode:</p>
<ul>
    <li>Template Mode now automatically maps your Excel columns to template paragraphs</li>
    <li>Works with any column letters (A, B, C or custom names)</li>
    <li>Empty paragraphs in your template are preserved as spacers</li>
    <li>No longer limited to specific column configurations</li>
</ul>
<div class="tip-box">
    <strong>v7.0 Enhancement</strong>
    Dynamic column mapping - your columns are automatically matched to non-empty paragraphs in your template, regardless of which column letters you use. This means you can use columns B and C, or D and E, or any combination that fits your data structure.
</div>
<div class="info-box">
    <strong>Template Mode Settings</strong>
    When using Template Mode:<br>
    - Upload a .pptx template with your desired slide design<br>
    - Specify the Image Placeholder Name (shape containing the image)<br>
    - Specify the Text Placeholder Name (shape containing text)<br>
    - StimuPop preserves all template formatting (fonts, sizes, colors)<br>
    - Your columns are dynamically mapped to non-empty template paragraphs
</div>

<h3>Image Sizing</h3>
<p>Control how images are sized uniformly across all slides. These settings work in both Blank and Template modes:</p>
<table>
    <tr><th>Size Mode</th><th>Description</th></tr>
    <tr><td><strong>Fit to Box (Recommended)</strong></td><td>Images scale to fit within Max Width and Max Height while maintaining aspect ratio. All images appear uniform.</td></tr>
    <tr><td><strong>Fit Width Only</strong></td><td>Images have fixed width, height adjusts automatically based on aspect ratio.</td></tr>
    <tr><td><strong>Fit Height Only</strong></td><td>Images have fixed height, width adjusts automatically based on aspect ratio.</td></tr>
    <tr><td><strong>Stretch to Exact Size</strong></td><td>Images are forced to exact dimensions. May distort images.</td></tr>
</table>
<p>Image dimension sliders:</p>
<ul>
    <li><strong>Max Width:</strong> 0.0 - 9.0 inches (0 allows thumbnail-sized images)</li>
    <li><strong>Max Height:</strong> 0.0 - 7.0 inches (0 allows thumbnail-sized images)</li>
</ul>
<div class="tip-box">
    <strong>Tip</strong>
    Use 'Fit to Box' mode for product catalogs and photo albums where you want all images to appear the same size regardless of their original dimensions.
</div>

<h3>Layout Position (Blank Mode Only)</h3>
<p>These settings control slide layout and are only available when using Blank slide mode (not Template mode):</p>
<ul>
    <li><strong>Image Top Position:</strong> Distance from top of slide to image (Default: 0.5 inches)</li>
    <li><strong>Text Top Position:</strong> Distance from top of slide to text area, 0.0-9.0 inches (Default: 5.0 inches)</li>
    <li><strong>Slide Orientation:</strong> Portrait (tall) or Landscape (wide) (Default: Portrait)</li>
</ul>

<h3>Image Alignment (Blank Mode Only)</h3>
<p>Control how images are positioned within their bounding box. These settings only apply to Blank mode:</p>
<table>
    <tr><th>Alignment Type</th><th>Options</th></tr>
    <tr><td><strong>Vertical Alignment</strong></td><td>Top, Center (default), Bottom</td></tr>
    <tr><td><strong>Horizontal Alignment</strong></td><td>Left, Center (default), Right</td></tr>
</table>
<div class="tip-box">
    <strong>Tip</strong>
    Use 'Bottom' vertical alignment for variety cards where you want images anchored to the bottom of the image area, regardless of image height.
</div>

<h3>Text Spacing</h3>
<p>Control spacing and overflow behavior for text content:</p>
<ul>
    <li><strong>Paragraph Spacing:</strong> Space after each paragraph (0-24pt). Default 0pt means no extra spacing between lines.</li>
    <li><strong>Text Overflow:</strong> How text boxes handle content that exceeds the box size:
        <ul>
            <li><strong>Resize shape to fit text (default):</strong> The text box expands to fit all content. Text remains at original size.</li>
            <li><strong>Shrink text on overflow:</strong> Text automatically shrinks to fit within the text box boundaries.</li>
        </ul>
    </li>
</ul>
<div class="info-box">
    <strong>Text Overflow</strong>
    The Text Overflow dropdown gives you control over how StimuPop handles long text that exceeds the text box. 'Resize shape to fit text' keeps text readable but may extend beyond the slide. 'Shrink text on overflow' keeps text within bounds but may make it smaller.
</div>

<h3>Advanced Positioning</h3>
<p>For precise control over text column placement, enable 'Advanced Positioning' by checking the checkbox in Advanced Settings. This reveals per-column position controls:</p>
<ul>
    <li><strong>Auto (default):</strong> Text flows sequentially after the previous column. Position depends on content length of earlier columns.</li>
    <li><strong>Fixed:</strong> Text is placed at an exact position on the slide, regardless of other content. Creates a separate text box.</li>
</ul>
<p>Default fixed positions (useful for variety cards):</p>
<ul>
    <li><strong>Column E:</strong> 5.0 inches from top</li>
    <li><strong>Column F:</strong> 6.5 inches from top</li>
</ul>
<div class="tip-box">
    <strong>Tip</strong>
    Use Fixed positioning for columns E and F when you want them to appear in the same location on every slide, even if the Brand (C) and Product Name (D) vary in length.
</div>

<h3>Per-Column Formatting (Blank Mode Only)</h3>
<p>When using Blank mode, you can customize each text column individually in Advanced Settings:</p>
<ul>
    <li>Font size (8-48pt)</li>
    <li>Font family (Calibri, Arial, Times New Roman, etc.)</li>
    <li>Text color (color picker)</li>
    <li>Bold and Italic styles</li>
</ul>
<div class="tip-box">
    <strong>Tip</strong>
    Use larger, bold fonts for titles (Column C) and smaller regular fonts for descriptions (Column D) to create visual hierarchy.
</div>

<h1 id="troubleshooting">6. Troubleshooting</h1>

<h3>Application won't start</h3>
<ul>
    <li>Ensure you extracted all files from the ZIP</li>
    <li>Run as Administrator (right-click -> Run as administrator)</li>
    <li>Check that antivirus isn't blocking the application</li>
</ul>

<h3>Browser doesn't open automatically</h3>
<ul>
    <li>Wait for 'Server ready!' message in the terminal window</li>
    <li>Hold Ctrl and click the localhost link, or copy/paste the URL</li>
    <li>Manually open your browser and go to: http://localhost:8501</li>
    <li>Try a different browser (Chrome recommended)</li>
</ul>

<h3>'Column not found' error</h3>
<ul>
    <li>Check that your column letter/name matches exactly</li>
    <li>Column letters are case-insensitive (B = b)</li>
    <li>Column names ARE case-sensitive</li>
</ul>

<h3>Images not appearing on slides</h3>
<ul>
    <li>Ensure images are embedded in cells, not floating</li>
    <li>Check file paths are correct and accessible</li>
    <li>Verify image files are JPG, PNG, or GIF format</li>
</ul>

<h3>Presentation generation is slow</h3>
<ul>
    <li>Large images take longer to process</li>
    <li>Consider resizing images in Excel before upload</li>
    <li>Reduce the number of slides if testing</li>
</ul>

<h3>Text formatting not applied</h3>
<ul>
    <li>Ensure you configured formatting in Advanced Settings</li>
    <li>Check that column letters in formatting match your data</li>
</ul>

<h3>Images are different sizes on slides</h3>
<ul>
    <li>Open Advanced Settings and set Size Mode to 'Fit to Box'</li>
    <li>Adjust Max Width and Max Height to your desired dimensions</li>
    <li>All images will now scale uniformly</li>
</ul>

<h1 id="quick-reference">7. Quick Reference</h1>

<h2>Column Reference Cheat Sheet</h2>
<table>
    <tr><th>Reference Type</th><th>Example</th></tr>
    <tr><td>Single letter</td><td><code>B</code></td></tr>
    <tr><td>Multiple letters</td><td><code>C,D,E,F</code></td></tr>
    <tr><td>Column name</td><td><code>Title</code></td></tr>
    <tr><td>Mixed</td><td><code>B,Title,Description</code></td></tr>
    <tr><td>With spaces</td><td><code>"Product Name"</code></td></tr>
</table>

<h2>Keyboard Shortcuts</h2>
<table>
    <tr><th>Shortcut</th><th>Action</th></tr>
    <tr><td><code>Ctrl + C</code></td><td>Stop the application (in command window)</td></tr>
    <tr><td><code>F5</code></td><td>Refresh the browser page</td></tr>
    <tr><td><code>Ctrl + S</code></td><td>Save downloaded file (in browser)</td></tr>
</table>

<div class="info-box">
    <strong>Need Help?</strong>
    If you encounter issues not covered in this guide, please contact your system administrator or the development team with:<br>
    - A description of the problem<br>
    - The exact error message (if any)<br>
    - Your Excel file (if possible)
</div>

<div style="text-align: center; margin-top: 60px; padding: 40px;">
    <p style="font-size: 1.3em; color: var(--primary-blue); font-weight: bold;">Thank you for using StimuPop!</p>
    <p style="color: var(--text-gray);">Version 7.0.0</p>
</div>

</body>
</html>'''

    with open('StimuPop_User_Guide.html', 'w', encoding='utf-8') as f:
        f.write(html_content)
    print("[OK] HTML User Guide created successfully: StimuPop_User_Guide.html")


if __name__ == "__main__":
    create_user_guide()
    create_html_user_guide()
